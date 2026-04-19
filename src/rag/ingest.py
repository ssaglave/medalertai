"""
src/rag/ingest.py

Phase 3 - Source document collection and conversion for the RAG pipeline.
Owner: Greeshma (C1)

Responsibilities covered here:
  - Maintain a source registry for PA DOH EMS protocols, NFPA 1221 reference
    metadata, and WPRDC glossary/data-dictionary pages.
  - Convert local PDF, DOCX, TXT, Markdown, and HTML files to clean text.
  - Split extracted text into JSONL chunks with citation-ready metadata.
  - Run basic quality checks so downstream embedding does not ingest empty or
    garbled chunks.

Usage from the repo root:
    python -m src.rag.ingest --write-source-list
    python -m src.rag.ingest --input-dir data/external/rag/source_documents
    python -m src.rag.ingest --download
"""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import logging
import re
import sys
import urllib.error
import urllib.parse
import urllib.request
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable, Optional

from config.settings import PROCESSED_DATA_DIR, PROJECT_ROOT


log = logging.getLogger("medalertai.rag.ingest")

SOURCE_DIR = PROJECT_ROOT / "data" / "external" / "rag" / "source_documents"
OUTPUT_DIR = PROCESSED_DATA_DIR / "rag"
CHUNKS_PATH = OUTPUT_DIR / "chunks.jsonl"
SOURCE_LIST_PATH = SOURCE_DIR / "source_registry.json"

CHUNK_SIZE = 900
CHUNK_OVERLAP = 150
MIN_CHUNK_CHARS = 120
MIN_ALPHA_RATIO = 0.50

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}


DEFAULT_SOURCES = [
    {
        "source_id": "pa_doh_ems_regulations",
        "title": "Pennsylvania EMS Regulations and Protocols",
        "owner": "Pennsylvania Department of Health",
        "url": "https://www.pa.gov/agencies/health/healthcare-and-public-health-professionals/ems/ems-regulations",
        "category": "pa_doh_ems",
        "license_note": "Public government web page; use as source index for current protocol PDFs.",
        "download": True,
    },
    {
        "source_id": "pa_doh_2023_bls_protocols",
        "title": "Final 2023 Pennsylvania BLS Protocols",
        "owner": "Pennsylvania Department of Health",
        "url": "https://www.pa.gov/content/dam/copapwp-pagov/en/health/documents/topics/documents/ems/2023v1-2%20PA%20BLS%20Protocols.pdf",
        "category": "pa_doh_ems",
        "license_note": "Public government protocol PDF.",
        "download": True,
    },
    {
        "source_id": "pa_doh_2023_als_protocols",
        "title": "Final 2023 Pennsylvania ALS Protocols",
        "owner": "Pennsylvania Department of Health",
        "url": "https://www.pa.gov/agencies/health/healthcare-and-public-health-professionals/ems/ems-regulations",
        "category": "pa_doh_ems",
        "license_note": "Protocol PDF is linked from the PA DOH EMS regulations page; add the downloaded PDF locally if the URL changes.",
        "download": False,
    },
    {
        "source_id": "nfpa_1221_2019_reference",
        "title": "NFPA 1221-2019 Standard Reference Page",
        "owner": "National Fire Protection Association",
        "url": "https://webstore.ansi.org/standards/nfpa/nfpa12212019",
        "category": "nfpa_communications",
        "license_note": "Do not commit NFPA standard PDFs. Use this public metadata page or a locally licensed copy only.",
        "download": False,
    },
    {
        "source_id": "wprdc_data_dictionaries",
        "title": "WPRDC Data Dictionaries",
        "owner": "Western Pennsylvania Regional Data Center",
        "url": "https://wiki.wprdc.org/wiki/Data_Dictionaries",
        "category": "wprdc_glossary",
        "license_note": "Public WPRDC wiki page explaining data dictionary fields.",
        "download": True,
    },
]


@dataclass(frozen=True)
class Document:
    source_id: str
    title: str
    path: Path
    text: str
    metadata: dict


class _VisibleTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._skip_depth = 0
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, Optional[str]]]) -> None:
        if tag.lower() in {"script", "style", "noscript", "svg"}:
            self._skip_depth += 1
        if tag.lower() in {"p", "br", "div", "li", "tr", "h1", "h2", "h3"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "noscript", "svg"} and self._skip_depth:
            self._skip_depth -= 1
        if tag.lower() in {"p", "li", "tr", "h1", "h2", "h3"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        return " ".join(self.parts)


def _slugify(value: str) -> str:
    value = value.lower().strip()
    value = re.sub(r"[^a-z0-9]+", "_", value)
    return value.strip("_") or "document"


def _clean_text(text: str) -> str:
    text = html.unescape(text)
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" *\n *", "\n", text)
    return text.strip()


def _quality_score(text: str) -> dict:
    if not text:
        return {"chars": 0, "alpha_ratio": 0.0, "passes": False}
    alpha = sum(char.isalpha() for char in text)
    readable = sum(char.isprintable() or char.isspace() for char in text)
    alpha_ratio = alpha / max(len(text), 1)
    readable_ratio = readable / max(len(text), 1)
    return {
        "chars": len(text),
        "alpha_ratio": round(alpha_ratio, 3),
        "readable_ratio": round(readable_ratio, 3),
        "passes": len(text) >= MIN_CHUNK_CHARS and alpha_ratio >= MIN_ALPHA_RATIO and readable_ratio >= 0.95,
    }


def write_source_registry(path: Path = SOURCE_LIST_PATH) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(DEFAULT_SOURCES, indent=2) + "\n", encoding="utf-8")
    log.info("Wrote source registry to %s", _display_path(path))


def download_default_sources(source_dir: Path = SOURCE_DIR, overwrite: bool = False) -> list[Path]:
    source_dir.mkdir(parents=True, exist_ok=True)
    downloaded: list[Path] = []

    for source in DEFAULT_SOURCES:
        if not source.get("download"):
            log.info("Skipping manual/licensed source: %s", source["title"])
            continue

        url = source["url"]
        suffix = Path(urllib.parse.urlparse(url).path).suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            suffix = ".html"

        target = source_dir / f"{source['source_id']}{suffix}"
        if target.exists() and not overwrite:
            log.info("Already exists, skipping: %s", _display_path(target))
            downloaded.append(target)
            continue

        log.info("Downloading %s", url)
        request = urllib.request.Request(url, headers={"User-Agent": "MedAlertAI student RAG ingestion"})
        try:
            with urllib.request.urlopen(request, timeout=30) as response:
                target.write_bytes(response.read())
        except (urllib.error.URLError, TimeoutError) as exc:
            log.warning("Could not download %s: %s", url, exc)
            continue
        downloaded.append(target)
        log.info("Saved %s", _display_path(target))

    write_source_registry(source_dir / "source_registry.json")
    return downloaded


def extract_pdf_text(path: Path) -> list[tuple[int, str]]:
    try:
        from PyPDF2 import PdfReader
    except ImportError as exc:
        raise RuntimeError("PyPDF2 is required for PDF ingestion. Install requirements.txt first.") from exc

    reader = PdfReader(str(path))
    pages: list[tuple[int, str]] = []
    for page_number, page in enumerate(reader.pages, start=1):
        pages.append((page_number, _clean_text(page.extract_text() or "")))
    return pages


def extract_docx_text(path: Path) -> list[tuple[int, str]]:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX ingestion. Install requirements.txt first.") from exc

    document = docx.Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return [(1, _clean_text("\n".join(paragraphs)))]


def extract_html_text(path: Path) -> list[tuple[int, str]]:
    parser = _VisibleTextParser()
    parser.feed(path.read_text(encoding="utf-8", errors="ignore"))
    return [(1, _clean_text(parser.text()))]


def extract_plain_text(path: Path) -> list[tuple[int, str]]:
    return [(1, _clean_text(path.read_text(encoding="utf-8", errors="ignore")))]


def extract_text(path: Path) -> list[tuple[int, str]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix in {".html", ".htm"}:
        return extract_html_text(path)
    if suffix in {".txt", ".md"}:
        return extract_plain_text(path)
    raise ValueError(f"Unsupported file type: {path}")


def _source_metadata_for_file(path: Path, registry: list[dict]) -> dict:
    stem = path.stem.lower()
    for source in registry:
        if source["source_id"].lower() == stem:
            return source
    return {
        "source_id": _slugify(path.stem),
        "title": path.stem.replace("_", " ").title(),
        "owner": "Local project source",
        "url": "",
        "category": "local_document",
        "license_note": "Local document supplied by project team.",
    }


def load_source_registry(path: Path = SOURCE_LIST_PATH) -> list[dict]:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return DEFAULT_SOURCES


def load_documents(input_dir: Path = SOURCE_DIR, registry_path: Path = SOURCE_LIST_PATH) -> list[Document]:
    if not input_dir.exists():
        raise FileNotFoundError(f"Source document folder does not exist: {input_dir}")

    registry = load_source_registry(registry_path)
    documents: list[Document] = []

    files = sorted(path for path in input_dir.rglob("*") if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS)
    if not files:
        log.warning("No supported source documents found in %s", _display_path(input_dir))
        return documents

    for path in files:
        metadata = _source_metadata_for_file(path, registry)
        log.info("Extracting %s", _display_path(path))
        page_text = extract_text(path)
        combined = "\n\n".join(text for _, text in page_text if text.strip())
        quality = _quality_score(combined)
        if not quality["passes"]:
            log.warning("Document quality check is weak for %s: %s", path.name, quality)
        documents.append(
            Document(
                source_id=metadata["source_id"],
                title=metadata["title"],
                path=path,
                text=combined,
                metadata={**metadata, "quality": quality, "pages_or_sections": len(page_text)},
            )
        )

    return documents


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    if overlap >= chunk_size:
        raise ValueError("overlap must be smaller than chunk_size")

    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        if not current:
            current = paragraph
        elif len(current) + len(paragraph) + 2 <= chunk_size:
            current = f"{current}\n\n{paragraph}"
        else:
            chunks.extend(_split_long_text(current, chunk_size, overlap))
            current = paragraph

    if current:
        chunks.extend(_split_long_text(current, chunk_size, overlap))

    return [chunk for chunk in chunks if _quality_score(chunk)["passes"]]


def _split_long_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            sentence_end = max(text.rfind(". ", start, end), text.rfind("\n", start, end))
            if sentence_end > start + MIN_CHUNK_CHARS:
                end = sentence_end + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


def build_chunks(documents: Iterable[Document]) -> list[dict]:
    rows: list[dict] = []
    for document in documents:
        for chunk_index, chunk in enumerate(chunk_text(document.text), start=1):
            chunk_id = hashlib.sha1(f"{document.source_id}:{chunk_index}:{chunk}".encode("utf-8")).hexdigest()[:16]
            rows.append(
                {
                    "chunk_id": chunk_id,
                    "source_id": document.source_id,
                    "title": document.title,
                    "chunk_index": chunk_index,
                    "text": chunk,
                    "metadata": {
                        **document.metadata,
                        "file_name": document.path.name,
                        "relative_path": str(_display_path(document.path)),
                        "chunk_chars": len(chunk),
                    },
                }
            )
    return rows


def write_chunks(chunks: list[dict], output_path: Path = CHUNKS_PATH) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8") as handle:
        for chunk in chunks:
            handle.write(json.dumps(chunk, ensure_ascii=False) + "\n")
    log.info("Wrote %d chunks to %s", len(chunks), _display_path(output_path))


def run_ingestion(
    input_dir: Path = SOURCE_DIR,
    output_path: Path = CHUNKS_PATH,
    registry_path: Path = SOURCE_LIST_PATH,
) -> list[dict]:
    documents = load_documents(input_dir=input_dir, registry_path=registry_path)
    chunks = build_chunks(documents)
    write_chunks(chunks, output_path)
    return chunks


def _display_path(path: Path) -> Path:
    try:
        return path.resolve().relative_to(PROJECT_ROOT)
    except ValueError:
        return path


def _parse_args(argv: Optional[list[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="MedAlertAI Phase 3 RAG source ingestion")
    parser.add_argument("--input-dir", type=Path, default=SOURCE_DIR)
    parser.add_argument("--output", type=Path, default=CHUNKS_PATH)
    parser.add_argument("--registry", type=Path, default=SOURCE_LIST_PATH)
    parser.add_argument("--download", action="store_true", help="Download public default sources before chunking.")
    parser.add_argument("--overwrite", action="store_true", help="Overwrite downloaded source files.")
    parser.add_argument("--write-source-list", action="store_true", help="Write the default source registry and exit.")
    return parser.parse_args(argv)


def main(argv: Optional[list[str]] = None) -> int:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s")
    args = _parse_args(argv)

    if args.write_source_list:
        write_source_registry(args.registry)
        return 0

    if args.download:
        download_default_sources(args.input_dir, overwrite=args.overwrite)

    chunks = run_ingestion(input_dir=args.input_dir, output_path=args.output, registry_path=args.registry)
    if not chunks:
        log.warning("No chunks were produced. Add source documents under %s or run with --download.", _display_path(args.input_dir))
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
